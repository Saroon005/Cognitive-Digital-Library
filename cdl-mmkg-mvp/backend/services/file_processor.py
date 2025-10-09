"""
Service for extracting text from various file formats
"""
import os
from typing import Optional


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text = ""
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        
        return text.strip()
    except ImportError:
        raise ImportError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Failed to extract PDF text: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except ImportError:
        raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Failed to extract DOCX text: {str(e)}")


def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Extract text from a file based on its extension
    Supports: .pdf, .docx, .txt
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif file_ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
